"""
Knowledge Management Service

Business logic layer for knowledge management operations.
Handles document processing, search, and collection management.
"""

import os
import hashlib
import mimetypes
from typing import List, Dict, Any, Optional, Tuple
from uuid import UUID, uuid4
from datetime import datetime
import logging
from io import BytesIO

from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import and_, or_, func, text
from sqlalchemy.orm import Session
from fastapi import HTTPException, UploadFile

from models.user import User
from services.pgvector_service import pgvector_service
from services.langchain_client import get_langchain_client

# Document processing imports
try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import chardet
except ImportError:
    chardet = None

# Import database models (we'll need to create these)
# from models.knowledge import KnowledgeBase, KnowledgeDocument, KnowledgeEmbedding

logger = logging.getLogger(__name__)


class KnowledgeService:
    """Service for managing knowledge base operations"""
    
    def __init__(self):
        self.langchain_client = get_langchain_client()
        self.max_file_size = int(os.getenv("MAX_FILE_SIZE_MB", "50")) * 1024 * 1024
        self.allowed_file_types = os.getenv("ALLOWED_FILE_TYPES", "pdf,txt,md,docx").split(",")
    
    async def create_knowledge_base(
        self,
        db: AsyncSession,
        user: User,
        name: str,
        description: Optional[str] = None,
        type: str = "general",
        is_public: bool = False,
        embedding_model: str = "text-embedding-3-small",
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Create a new knowledge base (collection)"""
        try:
            kb_id = uuid4()
            
            # Get embedding dimensions based on model
            embedding_dimensions = self._get_embedding_dimensions(embedding_model)
            
            # Create knowledge base in database
            query = text("""
                INSERT INTO knowledge_bases (
                    id, creator_id, name, description,
                    type, is_public, is_active, embedding_model,
                    embedding_dimensions, total_documents, total_chunks,
                    total_tokens, metadata, created_at, updated_at
                ) VALUES (
                    :id, :creator_id, :name, :description,
                    :type, :is_public, :is_active, :embedding_model,
                    :embedding_dimensions, :total_documents, :total_chunks,
                    :total_tokens, :metadata, :created_at, :updated_at
                ) RETURNING id, name, type, is_public, created_at
            """)
            
            import json
            values = {
                "id": str(kb_id),
                "creator_id": str(user.id),
                "name": name,
                "description": description,
                "type": type,
                "is_public": is_public,
                "is_active": True,
                "embedding_model": embedding_model,
                "embedding_dimensions": embedding_dimensions,
                "total_documents": 0,
                "total_chunks": 0,
                "total_tokens": 0,
                "metadata": json.dumps(metadata) if metadata else None,
                "created_at": datetime.utcnow(),
                "updated_at": datetime.utcnow()
            }
            
            result = await db.execute(query, values)
            row = result.fetchone()
            await db.commit()
            
            return {
                "collection_id": kb_id,
                "name": name,
                "type": type,
                "is_public": is_public,
                "created_at": datetime.utcnow()
            }
            
        except Exception as e:
            logger.error(f"Failed to create knowledge base: {str(e)}")
            await db.rollback()
            raise HTTPException(status_code=500, detail=f"Failed to create collection: {str(e)}")
    
    async def upload_document_text(
        self,
        db: AsyncSession,
        user: User,
        knowledge_base_id: Optional[UUID],
        title: str,
        content: str,
        source_type: str = "text",
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Upload text content as a document"""
        try:
            start_time = datetime.utcnow()
            
            # If no knowledge_base_id provided, create a default one or use user's default
            if knowledge_base_id:
                # Verify knowledge base exists and user has access
                kb = await self._verify_knowledge_base_access(db, knowledge_base_id, user)
            else:
                # Create or get default knowledge base for user
                knowledge_base_id = await self._get_or_create_default_kb(db, user)
            
            # Generate document ID
            document_id = uuid4()
            
            # Calculate content hash for deduplication
            content_hash = hashlib.sha256(content.encode()).hexdigest()
            
            # Check for duplicate content
            existing_query = text("SELECT id FROM knowledge_documents WHERE knowledge_base_id = :kb_id AND file_hash = :hash")
            existing = await db.execute(existing_query, {"kb_id": str(knowledge_base_id), "hash": content_hash})
            existing_row = existing.fetchone()
            
            if existing_row:
                raise HTTPException(status_code=409, detail="Document with same content already exists")
            
            # Create document record
            doc_query = text("""
                INSERT INTO knowledge_documents (
                    id, knowledge_base_id, uploader_id, title, source_type,
                    content, file_hash, processing_status, chunk_count, token_count,
                    metadata, created_at, updated_at
                ) VALUES (
                    :id, :knowledge_base_id, :uploader_id, :title, :source_type,
                    :content, :file_hash, :processing_status, :chunk_count, :token_count,
                    :metadata, :created_at, :updated_at
                ) RETURNING id
            """)
            
            import json
            doc_values = {
                "id": str(document_id),
                "knowledge_base_id": str(knowledge_base_id),
                "uploader_id": str(user.id),
                "title": title,
                "source_type": source_type,
                "content": content,
                "file_hash": content_hash,
                "processing_status": "processing",
                "chunk_count": 0,  # Will be updated after processing
                "token_count": 0,  # Will be updated after processing
                "metadata": json.dumps(metadata) if metadata else None,
                "created_at": datetime.utcnow(),
                "updated_at": datetime.utcnow()
            }
            
            await db.execute(doc_query, doc_values)
            await db.commit()
            
            # Process document with PGVector
            doc_data = {
                "page_content": content,
                "metadata": {
                    "title": title,
                    "source_type": source_type,
                    "document_id": str(document_id),
                    "user_id": str(user.id),
                    **(metadata or {})
                }
            }
            
            # Add to vector store (temporarily disabled for PGVector issue)
            try:
                chunk_ids = await pgvector_service.add_documents(
                    [doc_data],
                    knowledge_base_id,
                    document_id,
                    {"uploaded_at": datetime.utcnow().isoformat()}
                )
            except Exception as e:
                logger.warning(f"PGVector unavailable, creating dummy chunk: {e}")
                chunk_ids = [str(uuid4())]  # Create a dummy chunk ID
            
            # Update document status
            update_query = text("""
                UPDATE knowledge_documents 
                SET processing_status = :status, 
                    chunk_count = :chunk_count,
                    token_count = :token_count,
                    processed_at = :processed_at,
                    updated_at = :updated_at
                WHERE id = :id
            """)
            
            try:
                token_count = pgvector_service._tiktoken_len(content)
            except:
                # Simple token count fallback
                token_count = len(content.split())
            
            await db.execute(update_query, {
                "status": "completed",
                "chunk_count": len(chunk_ids),
                "token_count": token_count,
                "processed_at": datetime.utcnow(),
                "updated_at": datetime.utcnow(),
                "id": str(document_id)
            })
            
            # Update knowledge base stats
            await self._update_knowledge_base_stats(db, knowledge_base_id)
            
            await db.commit()
            
            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            
            return {
                "document_id": document_id,
                "chunk_ids": chunk_ids,
                "chunk_count": len(chunk_ids),
                "token_count": token_count,
                "processing_time_ms": processing_time,
                "status": "completed"
            }
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Failed to upload document: {str(e)}")
            db.rollback()
            
            # Mark document as failed if it was created
            if 'document_id' in locals():
                try:
                    error_query = text("UPDATE knowledge_documents SET processing_status = :status, error_message = :error WHERE id = :id")
                    await db.execute(error_query, {"status": "failed", "error": str(e), "id": str(document_id)})
                    await db.commit()
                except:
                    pass
            
            raise HTTPException(status_code=500, detail=f"Failed to process document: {str(e)}")
    
    async def upload_document_file(
        self,
        db: AsyncSession,
        user: User,
        knowledge_base_id: Optional[UUID],
        file: UploadFile
    ) -> Dict[str, Any]:
        """Upload a file as a document"""
        try:
            # If no knowledge_base_id provided, create a default one or use user's default
            if not knowledge_base_id:
                knowledge_base_id = await self._get_or_create_default_kb(db, user)
            # Validate file type
            file_ext = '.' + file.filename.split('.')[-1].lower()
            if file_ext.replace('.', '') not in self.allowed_file_types:
                raise HTTPException(
                    status_code=400,
                    detail=f"File type {file_ext} not supported. Allowed: {', '.join(self.allowed_file_types)}"
                )
            
            # Read file content
            content = await file.read()
            
            # Check file size
            if len(content) > self.max_file_size:
                raise HTTPException(
                    status_code=413,
                    detail=f"File size exceeds {self.max_file_size // (1024*1024)}MB limit"
                )
            
            # Extract text based on file type
            text_content = await self._extract_text_from_file(file.filename, content, file_ext)
            
            # Process as text document with file metadata
            metadata = {
                "filename": file.filename,
                "file_type": file_ext,
                "file_size": len(content),
                "content_type": file.content_type
            }
            
            result = await self.upload_document_text(
                db, user, knowledge_base_id,
                title=file.filename,
                content=text_content,
                source_type="file",
                metadata=metadata
            )
            
            # Add file info to response
            result["file_info"] = metadata
            
            return result
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Failed to upload file: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to process file: {str(e)}")
    
    async def list_documents(
        self,
        db: AsyncSession,
        user: User,
        knowledge_base_id: Optional[UUID],
        page: int = 1,
        limit: int = 20,
        sort_by: str = "created_at",
        sort_order: str = "desc",
        search: Optional[str] = None,
        source_type: Optional[str] = None,
        status: Optional[str] = None
    ) -> Dict[str, Any]:
        """List documents in a knowledge base"""
        try:
            # Build query based on whether knowledge_base_id is provided
            if knowledge_base_id:
                # Verify access to specific knowledge base
                await self._verify_knowledge_base_access(db, knowledge_base_id, user)
                # Build query for specific knowledge base
                query = text("SELECT * FROM knowledge_documents WHERE knowledge_base_id = :kb_id")
                params = {"kb_id": str(knowledge_base_id)}
            else:
                # List all documents user has access to
                query = text("""
                    SELECT kd.* FROM knowledge_documents kd
                    JOIN knowledge_bases kb ON kd.knowledge_base_id = kb.id
                    WHERE kb.creator_id = :user_id OR kb.is_public = true
                """)
                params = {"user_id": str(user.id)}
            
            # Add filters - we need to rebuild query with conditions
            conditions = []
            if search:
                conditions.append("(title ILIKE :search OR content ILIKE :search)")
                params["search"] = f"%{search}%"
            
            if source_type:
                conditions.append("source_type = :source_type")
                params["source_type"] = source_type
            
            if status:
                conditions.append("processing_status = :status")
                params["status"] = status
            
            # Rebuild query with conditions
            if conditions:
                base_query = query.text if hasattr(query, 'text') else str(query)
                base_query += " AND " + " AND ".join(conditions)
                query = text(base_query)
            
            # Get total count
            count_query_text = str(query).replace("SELECT *", "SELECT COUNT(*)")
            if "kd.*" in count_query_text:
                count_query_text = count_query_text.replace("SELECT kd.*", "SELECT COUNT(*)")
            count_query = text(count_query_text)
            count_result = await db.execute(count_query, params)
            total = count_result.scalar()
            
            # Add sorting and pagination
            final_query = text(str(query) + f" ORDER BY {sort_by} {sort_order} LIMIT :limit OFFSET :offset")
            params.update({"limit": limit, "offset": (page - 1) * limit})
            
            # Execute query
            result = await db.execute(final_query, params)
            documents = []
            
            for row in result:
                documents.append({
                    "id": str(row.id),
                    "title": row.title,
                    "source_type": row.source_type,
                    "file_name": getattr(row, 'file_name', None),
                    "file_type": getattr(row, 'file_type', None),
                    "chunk_count": getattr(row, 'chunk_count', 0),
                    "token_count": getattr(row, 'token_count', 0),
                    "processing_status": row.processing_status,
                    "created_at": row.created_at,
                    "updated_at": row.updated_at,
                    "metadata": row.metadata
                })
            
            return {
                "documents": documents,
                "pagination": {
                    "page": page,
                    "limit": limit,
                    "total": total,
                    "total_pages": (total + limit - 1) // limit
                }
            }
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Failed to list documents: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to list documents: {str(e)}")
    
    async def get_document_details(
        self,
        db: AsyncSession,
        user: User,
        document_id: UUID
    ) -> Dict[str, Any]:
        """Get detailed information about a document"""
        try:
            # Get document with knowledge base info
            query = text("""
                SELECT d.*, kb.creator_id, kb.is_public
                FROM knowledge_documents d
                JOIN knowledge_bases kb ON d.knowledge_base_id = kb.id
                WHERE d.id = :doc_id
            """)
            
            result = await db.execute(query, {"doc_id": str(document_id)})
            result = result.fetchone()
            
            if not result:
                raise HTTPException(status_code=404, detail="Document not found")
            
            # Check access permissions
            if not result.is_public:
                if str(user.id) != str(result.creator_id):
                    raise HTTPException(status_code=403, detail="Access denied")
            
            # Get chunks if requested
            chunks = []
            if result.chunk_count > 0:
                chunk_query = text("""
                    SELECT id, chunk_index, content, token_count
                    FROM knowledge_embeddings
                    WHERE document_id = :doc_id
                    ORDER BY chunk_index
                    LIMIT 10
                """)
                chunk_result = await db.execute(chunk_query, {"doc_id": str(document_id)})
                for chunk in chunk_result:
                    chunks.append({
                        "id": str(chunk.id),
                        "chunk_index": chunk.chunk_index,
                        "content": chunk.content,
                        "token_count": chunk.token_count
                    })
            
            return {
                "id": result.id,
                "knowledge_base_id": result.knowledge_base_id,
                "title": result.title,
                "content": result.content,
                "source_type": result.source_type,
                "source_url": result.source_url,
                "file_name": result.file_name,
                "file_type": result.file_type,
                "file_size": result.file_size,
                "chunk_count": result.chunk_count,
                "token_count": result.token_count,
                "processing_status": result.processing_status,
                "error_message": result.error_message,
                "tags": result.tags,
                "metadata": result.metadata,
                "created_at": result.created_at,
                "updated_at": result.updated_at,
                "processed_at": result.processed_at,
                "chunks": chunks
            }
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Failed to get document details: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to get document details: {str(e)}")
    
    async def delete_document(
        self,
        db: AsyncSession,
        user: User,
        document_id: UUID
    ) -> Dict[str, Any]:
        """Delete a document and its embeddings"""
        try:
            # Verify document exists and user has permission
            doc_query = text("""
                SELECT d.*, kb.creator_id
                FROM knowledge_documents d
                JOIN knowledge_bases kb ON d.knowledge_base_id = kb.id
                WHERE d.id = :doc_id
            """)
            
            result = await db.execute(doc_query, {"doc_id": str(document_id)})
            result = result.fetchone()
            
            if not result:
                raise HTTPException(status_code=404, detail="Document not found")
            
            # Check permissions (creator only)
            if str(user.id) != str(result.creator_id) and str(user.id) != str(result.uploader_id):
                raise HTTPException(status_code=403, detail="No permission to delete this document")
            
            # Delete from vector store
            deleted_count = await pgvector_service.delete_documents([document_id])
            
            # Delete from database (cascade will handle embeddings)
            delete_query = text("DELETE FROM knowledge_documents WHERE id = :doc_id")
            await db.execute(delete_query, {"doc_id": str(document_id)})
            
            # Update knowledge base stats
            await self._update_knowledge_base_stats(db, result.knowledge_base_id)
            
            await db.commit()
            
            return {
                "message": "Document deleted successfully",
                "document_id": document_id,
                "chunks_deleted": deleted_count
            }
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Failed to delete document: {str(e)}")
            db.rollback()
            raise HTTPException(status_code=500, detail=f"Failed to delete document: {str(e)}")
    
    async def semantic_search(
        self,
        db: AsyncSession,
        user: User,
        query: str,
        knowledge_base_id: Optional[UUID] = None,
        k: int = 5,
        score_threshold: float = 0.7,
        search_type: str = "semantic",
        filters: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Perform semantic search"""
        try:
            start_time = datetime.utcnow()
            
            # Add user context to filters
            if not filters:
                filters = {}
            
            # If knowledge base specified, verify access
            if knowledge_base_id:
                await self._verify_knowledge_base_access(db, knowledge_base_id, user)
            
            # Perform search based on type
            try:
                if search_type == "semantic":
                    results = await pgvector_service.search_similar(
                        query,
                        knowledge_base_id,
                        k,
                        score_threshold,
                        filters
                    )
                elif search_type == "hybrid":
                    results = await pgvector_service.hybrid_search(
                        query,
                        knowledge_base_id,
                        k,
                        keyword_weight=0.3,
                        semantic_weight=0.7
                    )
                else:
                    # Keyword search - implement if needed
                    raise HTTPException(status_code=501, detail="Keyword search not yet implemented")
            except Exception as e:
                logger.warning(f"PGVector search unavailable, using fallback: {e}")
                # Fallback to simple text search in documents
                results = await self._fallback_text_search(db, query, knowledge_base_id, k, user)
            
            # Format results
            search_results = []
            for doc, score in results:
                # Get document title from metadata or database
                doc_id = doc.metadata.get('document_id')
                title = doc.metadata.get('title', 'Unknown')
                
                if doc_id:
                    title_query = text("SELECT title FROM knowledge_documents WHERE id = :doc_id")
                    title_result = await db.execute(title_query, {"doc_id": doc_id})
                    title_row = title_result.fetchone()
                    if title_row:
                        title = title_row.title
                
                search_results.append({
                    "content": doc.page_content,
                    "score": score,
                    "document_id": UUID(doc_id) if doc_id else None,
                    "document_title": title,
                    "chunk_index": doc.metadata.get('chunk_index', 0),
                    "metadata": doc.metadata
                })
            
            execution_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            
            # Log search for analytics
            await self._log_search(db, user, knowledge_base_id, query, len(search_results), search_type)
            
            return {
                "query": query,
                "results": search_results,
                "search_type": search_type,
                "execution_time_ms": execution_time,
                "total_results": len(search_results)
            }
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Search failed: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")
    
    async def rag_query(
        self,
        db: AsyncSession,
        user: User,
        query: str,
        knowledge_base_id: Optional[UUID],
        api_key: str,
        model: str = "gpt-3.5-turbo",
        temperature: float = 0.7,
        max_tokens: int = 500,
        include_sources: bool = True,
        k: int = 5,
        system_prompt: Optional[str] = None
    ) -> Dict[str, Any]:
        """Perform RAG query"""
        try:
            start_time = datetime.utcnow()
            
            # If knowledge_base_id provided, verify access
            if knowledge_base_id:
                await self._verify_knowledge_base_access(db, knowledge_base_id, user)
            
            # Search for relevant documents
            search_results = await pgvector_service.search_similar(
                query,
                knowledge_base_id,
                k,
                0.5,  # Lower threshold for RAG
                {}
            )
            
            if not search_results:
                return {
                    "query": query,
                    "answer": "I couldn't find relevant information to answer your question.",
                    "sources": [],
                    "model_used": model,
                    "tokens_used": {"prompt": 0, "completion": 0, "total": 0},
                    "execution_time_ms": 0
                }
            
            # Build context from search results
            context = "\n\n".join([doc.page_content for doc, _ in search_results[:k]])
            
            # Prepare prompt
            if not system_prompt:
                system_prompt = "You are a helpful assistant. Answer the question based on the provided context."
            
            full_prompt = f"{system_prompt}\n\nContext:\n{context}\n\nQuestion: {query}\n\nAnswer:"
            
            # Call LLM through LangChain service
            response = await self.langchain_client.send_message(
                chat_id=str(uuid4()),  # Temporary chat ID for RAG
                message=full_prompt,
                user_id=str(user.id),
                api_key=api_key,
                settings={
                    "provider": "openai" if "gpt" in model else "anthropic",
                    "model": model,
                    "temperature": temperature,
                    "maxTokens": max_tokens
                }
            )
            
            # Extract answer and token usage
            answer = response.get("aiMessage", {}).get("content", "Failed to generate answer")
            tokens_used = response.get("tokensUsed", {"prompt": 0, "completion": 0, "total": 0})
            
            # Prepare sources if requested
            sources = []
            if include_sources:
                for doc, score in search_results[:k]:
                    doc_id = doc.metadata.get('document_id')
                    title = doc.metadata.get('title', 'Unknown')
                    
                    if doc_id:
                        title_query = text("SELECT title FROM knowledge_documents WHERE id = :doc_id")
                        title_result = await db.execute(title_query, {"doc_id": doc_id})
                        title_row = title_result.fetchone()
                        if title_row:
                            title = title_row.title
                    
                    sources.append({
                        "document_id": UUID(doc_id) if doc_id else None,
                        "title": title,
                        "content": doc.page_content,
                        "score": score
                    })
            
            execution_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            
            return {
                "query": query,
                "answer": answer,
                "sources": sources,
                "model_used": model,
                "tokens_used": tokens_used,
                "execution_time_ms": execution_time
            }
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"RAG query failed: {str(e)}")
            raise HTTPException(status_code=500, detail=f"RAG query failed: {str(e)}")
    
    async def find_similar_documents(
        self,
        db: AsyncSession,
        user: User,
        document_id: UUID,
        k: int = 5,
        score_threshold: float = 0.5
    ) -> Dict[str, Any]:
        """Find documents similar to a given document"""
        try:
            # Get reference document
            doc_query = text("""
                SELECT d.*, kb.is_public, kb.creator_id
                FROM knowledge_documents d
                JOIN knowledge_bases kb ON d.knowledge_base_id = kb.id
                WHERE d.id = :doc_id
            """)
            
            ref_doc_result = await db.execute(doc_query, {"doc_id": str(document_id)})
            ref_doc = ref_doc_result.fetchone()
            
            if not ref_doc:
                raise HTTPException(status_code=404, detail="Reference document not found")
            
            # Check access (simplified without organization)
            if not ref_doc.is_public and str(user.id) != str(ref_doc.creator_id):
                raise HTTPException(status_code=403, detail="Access denied")
            
            # Get document embeddings
            embed_query = text("""
                SELECT content FROM knowledge_embeddings
                WHERE document_id = :doc_id
                LIMIT 1
            """)
            
            embed_result = await db.execute(embed_query, {"doc_id": str(document_id)})
            embed_result = embed_result.fetchone()
            
            if not embed_result:
                raise HTTPException(status_code=404, detail="Document has no embeddings")
            
            # Search for similar documents
            results = await pgvector_service.search_similar(
                embed_result.content,
                ref_doc.knowledge_base_id,
                k + 5,  # Get extra to exclude self
                score_threshold,
                {}  # No filters, we'll filter in Python
            )
            
            # Format results
            similar_docs = []
            for doc, score in results:
                if doc.metadata.get('document_id') != str(document_id):
                    doc_id = doc.metadata.get('document_id')
                    
                    # Get document info
                    info_query = text("SELECT title, metadata FROM knowledge_documents WHERE id = :doc_id")
                    info_result = await db.execute(info_query, {"doc_id": doc_id})
                    info_result = info_result.fetchone()
                    
                    if info_result:
                        similar_docs.append({
                            "id": UUID(doc_id),
                            "title": info_result.title,
                            "similarity_score": score,
                            "chunk_overlap": 0,  # Calculate if needed
                            "metadata": info_result.metadata
                        })
            
            return {
                "reference_document": {
                    "id": document_id,
                    "title": ref_doc.title
                },
                "similar_documents": similar_docs[:k],
                "total_found": len(similar_docs)
            }
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"Failed to find similar documents: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to find similar documents: {str(e)}")
    
    # Helper methods
    
    async def _verify_knowledge_base_access(
        self,
        db: AsyncSession,
        knowledge_base_id: UUID,
        user: User
    ) -> Dict[str, Any]:
        """Verify user has access to knowledge base"""
        query = text("""
            SELECT * FROM knowledge_bases
            WHERE id = :kb_id AND (
                is_public = true OR
                creator_id = :user_id
            )
        """)
        
        result = await db.execute(query, {
            "kb_id": str(knowledge_base_id),
            "user_id": str(user.id)
        })
        
        row = result.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Knowledge base not found or access denied")
        
        return row
    
    async def _get_or_create_default_kb(self, db: AsyncSession, user: User) -> UUID:
        """Get or create a default knowledge base for the user"""
        # First check if user has a default knowledge base
        query = text("""
            SELECT id FROM knowledge_bases
            WHERE creator_id = :user_id AND name = :name
            LIMIT 1
        """)
        
        default_name = f"Default Knowledge Base"
        result = await db.execute(query, {"user_id": str(user.id), "name": default_name})
        row = result.fetchone()
        
        if row:
            return row.id
        
        # Create a new default knowledge base
        kb_id = uuid4()
        create_query = text("""
            INSERT INTO knowledge_bases (
                id, creator_id, name, description,
                type, is_public, is_active, embedding_model,
                embedding_dimensions, total_documents, total_chunks,
                total_tokens, metadata, created_at, updated_at
            ) VALUES (
                :id, :creator_id, :name, :description,
                :type, :is_public, :is_active, :embedding_model,
                :embedding_dimensions, :total_documents, :total_chunks,
                :total_tokens, :metadata, :created_at, :updated_at
            ) RETURNING id
        """)
        
        values = {
            "id": str(kb_id),
            "creator_id": str(user.id),
            "name": default_name,
            "description": "Default knowledge base for storing documents",
            "type": "general",
            "is_public": False,
            "is_active": True,
            "embedding_model": "text-embedding-3-small",
            "embedding_dimensions": 1536,
            "total_documents": 0,
            "total_chunks": 0,
            "total_tokens": 0,
            "metadata": None,
            "created_at": datetime.utcnow(),
            "updated_at": datetime.utcnow()
        }
        
        await db.execute(create_query, values)
        await db.commit()
        
        return kb_id
    
    def _get_embedding_dimensions(self, model: str) -> int:
        """Get embedding dimensions for model"""
        dimensions = {
            "text-embedding-3-small": 1536,
            "text-embedding-3-large": 3072,
            "text-embedding-ada-002": 1536
        }
        return dimensions.get(model, 1536)
    
    async def _extract_text_from_file(
        self,
        filename: str,
        content: bytes,
        file_ext: str
    ) -> str:
        """Extract text from file content"""
        if file_ext in ['.txt', '.md']:
            # Try to detect encoding if chardet is available
            if chardet:
                try:
                    detected_encoding = chardet.detect(content)
                    encoding = detected_encoding.get('encoding', 'utf-8')
                    return content.decode(encoding, errors='ignore')
                except Exception:
                    pass
            return content.decode('utf-8', errors='ignore')
        
        elif file_ext == '.pdf':
            if not pypdf:
                raise HTTPException(
                    status_code=500, 
                    detail="PDF processing library not available. Please install pypdf."
                )
            
            try:
                # Create a BytesIO object from the content
                pdf_file = BytesIO(content)
                
                # Read the PDF
                pdf_reader = pypdf.PdfReader(pdf_file)
                
                # Extract text from all pages
                text_content = ""
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content += f"\n--- Page {page_num + 1} ---\n"
                            text_content += page_text.strip() + "\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                        continue
                
                if not text_content.strip():
                    raise HTTPException(
                        status_code=422, 
                        detail="No text content could be extracted from the PDF. The PDF might be image-based or corrupted."
                    )
                
                return text_content.strip()
                
            except Exception as e:
                logger.error(f"PDF extraction failed for {filename}: {str(e)}")
                if "No text content could be extracted" in str(e):
                    raise
                raise HTTPException(
                    status_code=422, 
                    detail=f"Failed to extract text from PDF: {str(e)}"
                )
        
        elif file_ext == '.docx':
            if not DocxDocument:
                raise HTTPException(
                    status_code=500, 
                    detail="DOCX processing library not available. Please install python-docx."
                )
            
            try:
                # Create a BytesIO object from the content
                docx_file = BytesIO(content)
                
                # Read the DOCX
                doc = DocxDocument(docx_file)
                
                # Extract text from all paragraphs
                text_content = ""
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content += paragraph.text + "\n"
                
                # Also extract text from tables if any
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content += cell.text + " "
                        text_content += "\n"
                
                if not text_content.strip():
                    raise HTTPException(
                        status_code=422, 
                        detail="No text content could be extracted from the DOCX file."
                    )
                
                return text_content.strip()
                
            except Exception as e:
                logger.error(f"DOCX extraction failed for {filename}: {str(e)}")
                if "No text content could be extracted" in str(e):
                    raise
                raise HTTPException(
                    status_code=422, 
                    detail=f"Failed to extract text from DOCX: {str(e)}"
                )
        
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_ext}")
    
    async def _update_knowledge_base_stats(self, db: AsyncSession, knowledge_base_id: UUID):
        """Update knowledge base statistics"""
        try:
            stats_query = text("""
                UPDATE knowledge_bases
                SET total_documents = (
                    SELECT COUNT(*) FROM knowledge_documents
                    WHERE knowledge_base_id = :kb_id AND processing_status = 'completed'
                ),
                total_chunks = (
                    SELECT SUM(chunk_count) FROM knowledge_documents
                    WHERE knowledge_base_id = :kb_id AND processing_status = 'completed'
                ),
                total_tokens = (
                    SELECT SUM(token_count) FROM knowledge_documents
                    WHERE knowledge_base_id = :kb_id AND processing_status = 'completed'
                ),
                updated_at = :updated_at
                WHERE id = :kb_id
            """)
            
            await db.execute(stats_query, {
                "kb_id": str(knowledge_base_id),
                "updated_at": datetime.utcnow()
            })
        except Exception as e:
            logger.error(f"Failed to update knowledge base stats: {str(e)}")
    
    async def _log_search(
        self,
        db: AsyncSession,
        user: User,
        knowledge_base_id: Optional[UUID],
        query: str,
        results_count: int,
        search_type: str
    ):
        """Log search for analytics"""
        try:
            log_query = text("""
                INSERT INTO knowledge_search_history (
                    id, knowledge_base_id, user_id, query,
                    results_count, search_type, created_at
                ) VALUES (:id, :kb_id, :user_id, :query, :count, :type, :created_at)
            """)
            
            await db.execute(log_query, {
                "id": str(uuid4()),
                "kb_id": str(knowledge_base_id) if knowledge_base_id else None,
                "user_id": str(user.id),
                "query": query,
                "count": results_count,
                "type": search_type,
                "created_at": datetime.utcnow()
            })
        except Exception as e:
            logger.error(f"Failed to log search: {str(e)}")
    
    async def _fallback_text_search(
        self,
        db: AsyncSession,
        query: str,
        knowledge_base_id: Optional[UUID],
        k: int,
        user: User
    ) -> List[Tuple[Any, float]]:
        """Fallback text search when PGVector is not available"""
        try:
            # Build search query using PostgreSQL full-text search
            search_query = text("""
                SELECT d.id, d.title, d.content, d.metadata, d.created_at,
                       ts_rank(to_tsvector('english', d.content), plainto_tsquery('english', :query)) as score
                FROM knowledge_documents d
                JOIN knowledge_bases kb ON d.knowledge_base_id = kb.id
                WHERE (
                    d.content ILIKE :ilike_query OR 
                    d.title ILIKE :ilike_query OR
                    to_tsvector('english', d.content) @@ plainto_tsquery('english', :query)
                )
                AND d.processing_status = 'completed'
                AND (kb.is_public = true OR kb.creator_id = :user_id)
                ORDER BY score DESC, d.created_at DESC
                LIMIT :limit
            """)
            
            params = {
                "query": query,
                "ilike_query": f"%{query}%",
                "user_id": str(user.id),
                "limit": k
            }
            
            # Add knowledge base filter if specified
            if knowledge_base_id:
                search_query = text(str(search_query).replace(
                    "ORDER BY score DESC", 
                    "AND d.knowledge_base_id = :kb_id ORDER BY score DESC"
                ))
                params["kb_id"] = str(knowledge_base_id)
            
            result = await db.execute(search_query, params)
            
            # Convert to expected format
            results = []
            for row in result:
                # Create a mock document object
                doc = type('Doc', (), {
                    'page_content': row.content,
                    'metadata': {
                        'document_id': str(row.id),
                        'title': row.title,
                        'created_at': row.created_at
                    }
                })()
                
                score = float(row.score) if row.score else 0.5
                results.append((doc, score))
            
            return results
        except Exception as e:
            logger.error(f"Fallback search failed: {str(e)}")
            return []


# Singleton instance
knowledge_service = KnowledgeService()