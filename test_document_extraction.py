#!/usr/bin/env python3
"""
Test document extraction for PDF and DOCX files
"""

import requests
import json
from pathlib import Path

# Create test DOCX file
try:
    from docx import Document
    
    doc = Document()
    doc.add_heading('Test DOCX Document', 0)
    doc.add_paragraph('This is a test DOCX document for the Arketic knowledge base.')
    doc.add_heading('Features', level=1)
    doc.add_paragraph('The system supports multiple document formats:')
    doc.add_paragraph('• PDF files')
    doc.add_paragraph('• DOCX files')
    doc.add_paragraph('• TXT files')
    doc.add_paragraph('• Markdown files')
    
    # Add a table
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Feature'
    hdr_cells[1].text = 'Status'
    hdr_cells[2].text = 'Notes'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Text Extraction'
    row_cells[1].text = 'Working'
    row_cells[2].text = 'Extracts from all formats'
    
    doc.save('test_document.docx')
    print("✅ Created test_document.docx")
except ImportError:
    print("⚠️ python-docx not installed locally, skipping DOCX test")

# Create test PDF file
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    
    c = canvas.Canvas("test_document.pdf", pagesize=letter)
    c.drawString(100, 750, "Test PDF Document")
    c.drawString(100, 700, "This is a test PDF for the Arketic knowledge base.")
    c.drawString(100, 650, "Page 1 Content:")
    c.drawString(100, 620, "- Multi-file upload support")
    c.drawString(100, 590, "- Text extraction capabilities")
    c.drawString(100, 560, "- Embedding generation")
    
    c.showPage()
    c.drawString(100, 750, "Page 2")
    c.drawString(100, 700, "Additional content on the second page.")
    c.drawString(100, 650, "This demonstrates multi-page PDF support.")
    
    c.save()
    print("✅ Created test_document.pdf")
except ImportError:
    print("⚠️ reportlab not installed locally, skipping PDF test")

# Test upload with API
API_BASE_URL = "http://localhost:8000"

# Authenticate
auth_response = requests.post(
    f"{API_BASE_URL}/api/v1/auth/login",
    json={
        "email": "test@arketic.com",
        "password": "testpass123",
        "remember_me": False
    }
)

if auth_response.status_code == 200:
    token = auth_response.json()["access_token"]
    headers = {"Authorization": f"Bearer {token}"}
    
    print("\n📤 Testing individual file uploads...")
    
    # Test each file that exists
    for file_path in ["test_document.docx", "test_document.pdf"]:
        if Path(file_path).exists():
            with open(file_path, 'rb') as f:
                files = {'file': (file_path, f, 'application/octet-stream')}
                response = requests.post(
                    f"{API_BASE_URL}/api/v1/knowledge/upload/file",
                    files=files,
                    headers=headers
                )
                
                if response.status_code == 201:
                    result = response.json()
                    print(f"\n✅ {file_path} uploaded successfully!")
                    print(f"   - Document ID: {result.get('document_id')}")
                    print(f"   - Chunks: {result.get('chunk_count')}")
                    print(f"   - Tokens: {result.get('token_count')}")
                    print(f"   - Status: {result.get('status')}")
                else:
                    print(f"\n❌ Failed to upload {file_path}: {response.status_code}")
                    print(f"   Response: {response.text}")
    
    # Clean up
    for file_path in ["test_document.docx", "test_document.pdf"]:
        if Path(file_path).exists():
            Path(file_path).unlink()
            print(f"🧹 Removed {file_path}")
else:
    print(f"❌ Authentication failed: {auth_response.status_code}")